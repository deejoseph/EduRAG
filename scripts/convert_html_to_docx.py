"""
将高考作文HTML文件转换为DOCX格式，便于导入知识库
"""

from pathlib import Path
from bs4 import BeautifulSoup
import re

def html_to_text(html_file):
    """从HTML中提取正文文本"""
    with open(html_file, 'r', encoding='utf-8') as f:
        html = f.read()
    
    soup = BeautifulSoup(html, 'html.parser')
    
    # 尝试不同的内容选择器
    content_div = (
        soup.find('div', class_='article') or 
        soup.find('div', class_='content') or
        soup.find('div', id='article') or
        soup.find('div', class_='text')
    )
    
    if content_div:
        # 移除脚本和样式
        for script in content_div(['script', 'style']):
            script.decompose()
        
        # 提取段落
        paragraphs = content_div.find_all('p')
        if paragraphs:
            text = '\n\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])
        else:
            text = content_div.get_text(separator='\n', strip=True)
    else:
        # fallback: 提取body中的所有段落
        body = soup.find('body')
        if body:
            paragraphs = body.find_all('p')
            text = '\n\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])
        else:
            text = soup.get_text(separator='\n', strip=True)
    
    # 清理多余空行
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    text = text.strip()
    
    return text


def convert_htmls_to_docx():
    """批量转换HTML为DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("ERROR: python-docx not installed")
        print("Please run: pip install python-docx")
        return
    
    html_dir = Path("data/gaokao_essays/pdfs")
    docx_dir = Path("data/gaokao_essays/docx")
    docx_dir.mkdir(parents=True, exist_ok=True)
    
    html_files = list(html_dir.glob("*.html"))
    print(f"Found {len(html_files)} HTML files")
    
    converted = 0
    
    for html_file in html_files:
        print(f"\nProcessing: {html_file.name}")
        
        # 提取文本
        text = html_to_text(html_file)
        
        if not text or len(text) < 100:
            print(f"  SKIP: Content too short ({len(text)} chars)")
            continue
        
        print(f"  Extracted {len(text)} characters")
        
        # 创建DOCX文档
        doc = Document()
        
        # 添加标题（从文件名提取）
        title = html_file.stem.replace('_', ' ')
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分隔线
        doc.add_paragraph('_' * 50)
        
        # 添加正文
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para_format = para.paragraph_format
                para_format.first_line_indent = Pt(24)  # 首行缩进
                para_format.line_spacing = 1.5  # 1.5倍行距
                
                # 设置字体
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
        
        # 保存DOCX
        docx_file = docx_dir / f"{html_file.stem}.docx"
        doc.save(str(docx_file))
        
        print(f"  SAVED: {docx_file.name}")
        converted += 1
    
    print(f"\n{'='*60}")
    print(f"Conversion complete: {converted}/{len(html_files)} files")
    print(f"Output directory: {docx_dir.absolute()}")
    print(f"{'='*60}")


if __name__ == '__main__':
    convert_htmls_to_docx()
